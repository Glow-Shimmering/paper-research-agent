"""Deterministic DOCX renderer for research artifacts."""

from __future__ import annotations

import io
import re
import zipfile
from datetime import datetime, timezone

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from pragent.research import DEEP_READ_FIELD_LABELS

from .common import citation_output, parse_artifact_content
from .models import FrozenArtifactExport

_BLUE = RGBColor(0x2E, 0x74, 0xB5)
_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
_MUTED = RGBColor(0x66, 0x66, 0x66)
_XML_INVALID = re.compile(r"[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD]")


def render_docx(snapshot: FrozenArtifactExport) -> bytes:
    content = parse_artifact_content(snapshot)
    citations = citation_output(snapshot, content)
    labels = iter(citations.citations)
    document = Document()
    _configure_document(document, snapshot)
    _add_masthead(document, snapshot, content)
    artifact_type = snapshot.artifact.artifact_type
    if artifact_type == "deep_read":
        _render_deep_read(document, content, labels)
    elif artifact_type == "comparison":
        _render_comparison(document, snapshot, content, labels)
    elif artifact_type == "review_outline":
        _render_outline(document, content, labels)
    elif artifact_type == "review_section":
        _render_section(document, content, labels)
    _render_references(document, citations.bibliography)
    _render_evidence_appendix(document, snapshot)
    raw = io.BytesIO()
    document.save(raw)
    return _deterministic_package(raw.getvalue())


def _configure_document(document, snapshot) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _set_style_font(normal, "Calibri", 11, RGBColor(0, 0, 0))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, _BLUE, 16, 8),
        ("Heading 2", 13, _BLUE, 12, 6),
        ("Heading 3", 12, _DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        _set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    created = _docx_datetime(snapshot.revision.created_at)
    properties = document.core_properties
    properties.title = _clean(snapshot.artifact.title)
    properties.subject = "PRAgent research artifact export"
    properties.author = "PRAgent"
    properties.last_modified_by = "PRAgent"
    properties.created = created
    properties.modified = created
    properties.keywords = (
        f"artifact={snapshot.artifact.id};revision={snapshot.revision.revision_number}"
    )

    header = section.header.paragraphs[0]
    header.text = "PRAgent research export"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_runs(header, 9, _MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = f"{snapshot.project.title} | revision {snapshot.revision.revision_number}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_runs(footer, 9, _MUTED)


def _add_masthead(document, snapshot, content) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(_clean(snapshot.artifact.title or _content_title(content)))
    _set_run_font(run, "Calibri", 23, RGBColor(0, 0, 0), bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(_clean(snapshot.project.title))
    _set_run_font(run, "Calibri", 13, _MUTED)
    for label, value in (
        ("Artifact", snapshot.artifact.artifact_type),
        ("Revision", str(snapshot.revision.revision_number)),
        ("Citation style", snapshot.citation_style),
        ("Freshness", "stale" if snapshot.freshness.stale else "current"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, "Calibri", 10.5, RGBColor(0, 0, 0), bold=True)
        value_run = paragraph.add_run(_clean(value))
        _set_run_font(value_run, "Calibri", 10.5, RGBColor(0, 0, 0))


def _render_deep_read(document, card, labels) -> None:
    for name, field in card.ordered_fields():
        document.add_heading(DEEP_READ_FIELD_LABELS[name], level=1)
        text = "证据不足。" if field.insufficient_evidence else field.text
        _add_body(document, _with_citation(text, next(labels)))


def _render_comparison(document, snapshot, matrix, labels) -> None:
    document.add_heading("比较矩阵", level=1)
    table = document.add_table(rows=1, cols=3)
    _set_table_geometry(table, (1900, 2200, 5260))
    _set_header_row(table.rows[0], ("维度", "来源", "比较结论"))
    sources = snapshot.source_by_id()
    dimensions = {item.key: item for item in matrix.dimensions}
    for cell in matrix.cells:
        row = table.add_row()
        text = "证据不足" if cell.insufficient_evidence else cell.summary
        values = (
            dimensions[cell.dimension_key].label,
            sources[cell.source_id].title,
            _with_citation(text, next(labels)),
        )
        for target, value in zip(row.cells, values):
            target.text = _clean(value)
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _format_cell(target)
    _set_table_geometry(table, (1900, 2200, 5260))


def _render_outline(document, outline, labels) -> None:
    document.add_heading("研究问题", level=1)
    for question in outline.research_questions:
        _add_body(document, question.question)
    for section in outline.sections:
        document.add_heading(_clean(section.title), level=1)
        _add_body(document, section.objective)
        for claim in section.planned_claims:
            text = "证据不足" if claim.insufficient_evidence else claim.text
            _add_body(document, _with_citation(text, next(labels)))


def _render_section(document, section, labels) -> None:
    document.add_heading(_clean(section.section_title), level=1)
    for claim in section.claims:
        text = "证据不足" if claim.insufficient_evidence else claim.text
        _add_body(document, _with_citation(text, next(labels)))


def _render_references(document, bibliography) -> None:
    document.add_heading("参考文献", level=1)
    if not bibliography:
        _add_body(document, "无。")
        return
    for entry in bibliography:
        paragraph = _add_body(document, entry)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.left_indent = Inches(0.25)


def _render_evidence_appendix(document, snapshot) -> None:
    document.add_heading("Evidence appendix", level=1)
    if not snapshot.evidence:
        _add_body(document, "无。")
        return
    for item in snapshot.evidence:
        evidence = item.evidence
        document.add_heading(_clean(evidence.id), level=2)
        page = f", p. {evidence.page}" if evidence.page > 0 else ""
        _add_body(
            document,
            f"{evidence.title}{page}; field {item.link.field_path}",
            color=_MUTED,
        )
        paragraph = _add_body(document, evidence.text)
        paragraph.paragraph_format.left_indent = Inches(0.25)


def _add_body(document, text, *, color=RGBColor(0, 0, 0)):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(_clean(text))
    _set_run_font(run, "Calibri", 11, color)
    return paragraph


def _set_table_geometry(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = tbl_pr.find(qn("w:tblW"))
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths)))
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(value))
            _set_cell_margins(cell)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_header_row(row, labels) -> None:
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(repeat)
    for cell, label in zip(row.cells, labels):
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F2F4F7")
        cell._tc.get_or_add_tcPr().append(shading)
        _format_cell(cell, bold=True)


def _format_cell(cell, *, bold=False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.10
        for run in paragraph.runs:
            _set_run_font(run, "Calibri", 9.5, RGBColor(0, 0, 0), bold=bold)


def _set_style_font(style, name, size, color, *, bold=None) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def _set_run_font(run, name, size, color, *, bold=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _format_runs(paragraph, size, color) -> None:
    for run in paragraph.runs:
        _set_run_font(run, "Calibri", size, color)


def _deterministic_package(raw: bytes) -> bytes:
    source = zipfile.ZipFile(io.BytesIO(raw), "r")
    output = io.BytesIO()
    with source, zipfile.ZipFile(
        output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            target.writestr(info, source.read(name))
    return output.getvalue()


def _docx_datetime(value: str) -> datetime:
    try:
        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        return datetime(1980, 1, 1)
    if parsed.tzinfo is not None:
        parsed = parsed.astimezone(timezone.utc).replace(tzinfo=None)
    return parsed


def _content_title(content) -> str:
    return getattr(content, "title", getattr(content, "section_title", "研究导出"))


def _with_citation(text: str, citation: str) -> str:
    return f"{text} {citation}".rstrip()


def _clean(value) -> str:
    return _XML_INVALID.sub("", str(value))
