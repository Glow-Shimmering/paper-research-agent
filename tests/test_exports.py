import csv
import io
import json
import time
import zipfile
from pathlib import Path
from urllib.parse import quote

from docx import Document
from fastapi.testclient import TestClient

from pragent.exporting import (
    ArtifactExportService,
    ExportEnvelope,
    render_docx,
    render_json,
    render_markdown,
    safe_export_stem,
)
from pragent.models import Chunk, Paper
from pragent.research import (
    ComparisonCell,
    ComparisonDimension,
    ComparisonMatrix,
    ReviewOutline,
    ReviewSectionDraft,
)
from pragent.storage import JobRepository, ResearchRepository
from pragent.store import Store
from pragent.webapp import create_app


def _seed_comparison(tmp_path):
    database = tmp_path / "exports.db"
    store = Store(database)
    paper_ids = []
    evidence = []
    for index, title in enumerate(("第一篇 Agent 论文", "Second Agent Paper"), start=1):
        paper_id = store.upsert_paper(
            Paper(
                id=None,
                path=str(tmp_path / f"paper-{index}.pdf"),
                sha256=f"paper-sha-{index}",
                title=title,
                authors=["Alice Chen"] if index == 1 else ["Bob Smith"],
                year=2023 + index,
                page_count=1,
                has_text=True,
                indexed_at="2026-08-28T00:00:00+00:00",
            ),
            [Chunk(None, 0, 0, index, f"Exact evidence text {index}")],
        )
        paper_ids.append(paper_id)
        chunk = store.paper_chunks(paper_id)[0]
        evidence.append(store.pin_evidence(store.evidence_from_chunk(chunk.id)))

    repository = ResearchRepository(database)
    project = repository.create_project("可审计研究项目")
    sources = [repository.ensure_source_for_paper(item) for item in paper_ids]
    for position, source in enumerate(sources):
        repository.add_project_source(project.id, source.id, position=position)
    repository.add_source_record(
        sources[0].id,
        "fixture-provider",
        "fixture-record-1",
        {"venue": "Journal of Deterministic Exports"},
        retrieved_at="2026-08-28T00:00:00+00:00",
    )
    dimension = ComparisonDimension(key="method", label="方法")
    matrix = ComparisonMatrix(
        title="确定性比较矩阵",
        source_ids=[item.id for item in sources],
        dimensions=[dimension],
        cells=[
            ComparisonCell(
                source_id=source.id,
                dimension_key="method",
                summary=f"方法结论 {index}",
                evidence_refs=[
                    {"evidence_id": evidence[index - 1].id, "quote": evidence[index - 1].text}
                ],
            )
            for index, source in enumerate(sources, start=1)
        ],
    )
    artifact = repository.create_artifact(
        project.id,
        "comparison",
        title='CON: 比较矩阵? / "draft"',
        status="generating",
    )
    revision = repository.append_artifact_revision(
        artifact.id,
        matrix.model_dump(mode="json"),
        expected_artifact_version=artifact.version,
        created_by="system",
        evidence_links=[
            (
                item.id,
                f"$.cells.{sources[index].id}.method",
                0,
            )
            for index, item in enumerate(evidence)
        ],
        usage={"fixture": True},
        schema_version=1,
    )
    return store, repository, artifact.id, revision.id


def test_frozen_artifact_renderers_are_deterministic_and_schema_valid(tmp_path):
    store, repository, artifact_id, revision_id = _seed_comparison(tmp_path)
    service = ArtifactExportService(repository, store)
    snapshot = service.freeze_current(artifact_id)

    assert snapshot.revision.id == revision_id
    markdown = render_markdown(snapshot)
    assert markdown == render_markdown(snapshot)
    assert "比较矩阵" in markdown.decode("utf-8")
    assert "Evidence appendix" in markdown.decode("utf-8")

    raw_json = render_json(snapshot)
    assert raw_json == render_json(snapshot)
    envelope = ExportEnvelope.model_validate_json(raw_json)
    assert envelope.revision["id"] == revision_id
    assert envelope.sources[0]["provenance"][0]["provider"] == "fixture-provider"
    assert len(envelope.evidence) == 2

    docx = render_docx(snapshot)
    assert docx == render_docx(snapshot)
    with zipfile.ZipFile(io.BytesIO(docx)) as archive:
        assert "word/document.xml" in archive.namelist()
        assert {item.date_time for item in archive.infolist()} == {(1980, 1, 1, 0, 0, 0)}
    document = Document(io.BytesIO(docx))
    assert len(document.tables) == 1
    assert document.tables[0].rows[0].cells[0].text == "维度"
    assert any("参考文献" in paragraph.text for paragraph in document.paragraphs)

    repository.close()
    store.close()


def test_csv_round_trip_safe_names_and_atomic_current_revision_export(tmp_path):
    store, repository, artifact_id, revision_id = _seed_comparison(tmp_path)
    service = ArtifactExportService(repository, store)
    snapshot = service.freeze_current(artifact_id)
    rendered = service.render(snapshot, "csv")

    assert [item.suffix for item in rendered] == ["sources.csv", "comparison.csv"]
    sources_rows = list(csv.DictReader(io.StringIO(rendered[0].data.decode("utf-8"))))
    comparison_rows = list(
        csv.DictReader(io.StringIO(rendered[1].data.decode("utf-8")))
    )
    assert len(sources_rows) == 2
    assert len(comparison_rows) == 2
    assert sources_rows[0]["source_kind"] == "paper"
    assert isinstance(json.loads(sources_rows[0]["metadata_json"]), dict)

    output = tmp_path / "exports"
    first = service.export_current(artifact_id, "csv", output)
    second = service.export_current(artifact_id, "csv", output)
    assert [item.path for item in first] == [item.path for item in second]
    assert all(item.path.parent == output for item in first)
    assert all(item.path.is_file() and item.size == item.path.stat().st_size for item in first)
    assert not list(output.glob("*.tmp"))
    assert all(character not in first[0].path.name for character in '<>:"/\\|?*')
    assert revision_id in snapshot.revision.id
    assert safe_export_stem("CON", "artifact:unsafe", 2).startswith("research-artifact-")

    repository.close()
    store.close()


def test_review_export_assembles_current_section_drafts_and_marks_missing(tmp_path):
    store, repository, comparison_id, _revision_id = _seed_comparison(tmp_path)
    comparison = repository.get_artifact(comparison_id)
    project = repository.get_project(comparison.project_id)
    sources = [item.source for item in repository.list_project_sources(project.id).items]
    evidence = store.list_evidence()
    question = repository.create_question(project.id, "不同方法的证据边界是什么？")
    claims = [
        {
            "text": "两篇论文的方法边界可比较。",
            "source_ids": [item.id for item in sources],
            "evidence_refs": [
                {
                    "source_id": source.id,
                    "evidence_id": evidence[index].id,
                    "quote": evidence[index].text,
                }
                for index, source in enumerate(sources)
            ],
            "insufficient_evidence": False,
        }
    ]
    outline_content = ReviewOutline(
        title="完整综述",
        research_questions=[
            {"id": question.id, "question": question.question, "version": question.version}
        ],
        source_ids=[item.id for item in sources],
        comparison_artifact_id=comparison.id,
        comparison_revision_id=repository.get_current_artifact_revision(comparison.id).id,
        sections=[
            {
                "key": "methods",
                "title": "方法",
                "objective": "综合方法差异",
                "source_ids": [item.id for item in sources],
                "planned_claims": claims,
            },
            {
                "key": "limits",
                "title": "局限",
                "objective": "总结证据边界",
                "source_ids": [item.id for item in sources],
                "planned_claims": claims,
            },
        ],
    )
    outline = repository.create_artifact(
        project.id, "review_outline", title="完整综述", status="generating"
    )
    outline_revision = repository.append_artifact_revision(
        outline.id,
        outline_content.model_dump(mode="json"),
        expected_artifact_version=outline.version,
        created_by="system",
        evidence_links=[
            (item.id, f"$.sections.methods.claims.0", index)
            for index, item in enumerate(evidence)
        ],
        schema_version=1,
    )
    draft_content = ReviewSectionDraft(
        outline_artifact_id=outline.id,
        outline_revision_id=outline_revision.id,
        section_key="methods",
        section_title="方法",
        claims=[
            {
                "key": "method_synthesis",
                "text": "这是当前章节 revision 的整稿正文。",
                "citation_tokens": claims[0]["evidence_refs"],
                "insufficient_evidence": False,
            }
        ],
    )
    section = repository.create_artifact(
        project.id, "review_section", title="方法", status="generating"
    )
    repository.append_artifact_revision(
        section.id,
        draft_content.model_dump(mode="json"),
        expected_artifact_version=section.version,
        created_by="system",
        evidence_links=[
            (item.id, "$.claims.0", index) for index, item in enumerate(evidence)
        ],
        schema_version=1,
    )

    snapshot = ArtifactExportService(repository, store).freeze_current(outline.id)
    markdown = render_markdown(snapshot).decode("utf-8")
    assert len(snapshot.review_sections) == 1
    assert "这是当前章节 revision 的整稿正文" in markdown
    assert "本节尚未生成草稿" in markdown
    assert len(ExportEnvelope.model_validate_json(render_json(snapshot)).review_sections) == 1

    repository.close()
    store.close()


def test_web_export_preview_persistent_job_and_scoped_download(tmp_path):
    store, repository, artifact_id, revision_id = _seed_comparison(tmp_path)
    project_id = repository.get_artifact(artifact_id).project_id
    jobs = JobRepository(store.db_path)
    export_dir = tmp_path / "web-exports"
    app = create_app(
        store=store,
        research_repository=repository,
        job_repository=jobs,
        job_worker_count=1,
        export_directory=export_dir,
    )

    with TestClient(app, base_url="http://127.0.0.1") as client:
        page = client.get(
            f"/ui/projects/{project_id}/artifacts/{artifact_id}/exports"
        )
        assert page.status_code == 200
        assert "Markdown 预览" in page.text and "比较矩阵" in page.text
        preview = client.get(
            f"/api/v1/projects/{project_id}/artifacts/{artifact_id}/exports/preview"
        )
        assert preview.status_code == 200
        assert preview.json()["revision_id"] == revision_id
        invalid = client.post(
            f"/api/v1/projects/{project_id}/artifacts/{artifact_id}/exports",
            json={"format": "pdf"},
        )
        assert invalid.status_code == 400
        started = client.post(
            f"/api/v1/projects/{project_id}/artifacts/{artifact_id}/exports",
            json={"format": "json"},
        )
        assert started.status_code == 202
        job_id = started.json()["job"]["id"]
        deadline = time.monotonic() + 5
        while time.monotonic() < deadline:
            current = client.get(
                f"/api/v1/projects/{project_id}/artifacts/{artifact_id}/exports/jobs/{job_id}"
            )
            assert current.status_code == 200
            job = current.json()
            if job["terminal"]:
                break
            time.sleep(0.02)
        assert job["status"] == "succeeded"
        assert job["revision_id"] == revision_id
        filename = job["files"][0]["name"]
        downloaded = client.get(
            f"/api/v1/projects/{project_id}/artifacts/{artifact_id}/exports/jobs/"
            f"{job_id}/files/{quote(filename)}"
        )
        assert downloaded.status_code == 200
        assert ExportEnvelope.model_validate_json(downloaded.content).revision["id"] == revision_id
        missing = client.get(
            f"/api/v1/projects/{project_id}/artifacts/{artifact_id}/exports/jobs/"
            f"{job_id}/files/not-listed.json"
        )
        assert missing.status_code == 404

    jobs.close()
    repository.close()
    store.close()
